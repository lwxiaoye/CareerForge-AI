"""文件文本抽取共享 util。

PDF / DOCX / XLSX / TXT / MD / 图片 → 纯文本。
供 agent 附件、简历导入、面试官等多处复用。
"""
from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_pdf_text(path: Path, *, max_pages: int = 12, max_chars: int = 30000) -> str:
    """从 PDF 提取文本。pypdf 优先，pdfminer 兜底。"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        chunks: list[str] = []
        for index, page in enumerate(reader.pages[:max_pages], start=1):
            text = (page.extract_text() or "").strip()
            if text:
                chunks.append(f"[PDF 第 {index} 页]\n{text}")
        result = "\n\n".join(chunks)[:max_chars]
        if result:
            return result
    except Exception:
        pass

    try:
        from pdfminer.high_level import extract_text as _pdfminer_extract
        text = (_pdfminer_extract(str(path)) or "").strip()
        if text:
            return text[:max_chars]
    except Exception:
        pass

    return ""


def extract_docx_text(path: Path, *, max_chars: int = 30000) -> str:
    """从 DOCX 提取文本（段落 + 表格）。"""
    from docx import Document

    doc = Document(str(path))
    chunks = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables[:8]:
        for row in table.rows[:30]:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                chunks.append(" | ".join(values))
    return "\n".join(chunks)[:max_chars] or ""


def extract_xlsx_text(path: Path, *, max_chars: int = 12000) -> str:
    """从 XLSX 提取文本。"""
    from openpyxl import load_workbook

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    chunks: list[str] = []
    for sheet in workbook.worksheets[:5]:
        chunks.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(min_row=1, max_row=40, max_col=12, values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                chunks.append(" | ".join(values))
    return "\n".join(chunks)[:max_chars] or ""


def extract_image_summary(path: Path) -> str:
    """提取图片元信息（尺寸、色彩模式）。"""
    from PIL import Image

    with Image.open(path) as image:
        width, height = image.size
        mode = image.mode
    return f"图片附件已保存：{width}x{height}，色彩模式 {mode}。如所选模型支持视觉输入，将随请求一并传入。"


def extract_file_text(path: Path, content_type: str, ext: str, *, max_chars: int = 30000) -> str:
    """统一入口：根据扩展名/类型分派到对应抽取函数。ext 带不带前导点均可。"""
    ext = (ext or "").lower().lstrip(".")
    try:
        if ext == "pdf":
            return extract_pdf_text(path, max_chars=max_chars)
        if ext == "docx":
            return extract_docx_text(path, max_chars=max_chars)
        if ext in {"xlsx", "xls"}:
            return extract_xlsx_text(path, max_chars=max_chars)
        if ext in {"csv", "txt", "md", "json"}:
            return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
        if content_type.startswith("image/"):
            return extract_image_summary(path)
    except Exception as exc:
        logger.exception("文件文本抽取失败: %s", path)
        return ""
    return ""


def render_pdf_pages_to_png(
    path: Path,
    *,
    max_pages: int = 3,
    scale: float = 2.5,
) -> list[bytes]:
    """把 PDF 前 N 页渲染成 PNG 字节流，用于多模态 OCR 兜底。

    用 pypdfium2（PDFium 封装，跨平台、无系统依赖）。
    失败或无页时返回空列表。"""
    try:
        import pypdfium2 as pdfium
    except Exception as exc:
        logger.warning("pypdfium2 not available, skip render: %s", exc)
        return []

    try:
        document = pdfium.PdfDocument(str(path))
    except Exception as exc:
        logger.warning("pdfium open failed: %s", exc)
        return []

    total = len(document)
    if total == 0:
        return []
    page_count = min(total, max_pages)
    out: list[bytes] = []
    for index in range(page_count):
        try:
            page = document[index]
            pil_image = page.render(scale=scale).to_pil()
            from io import BytesIO
            buf = BytesIO()
            pil_image.save(buf, format="PNG", optimize=False)
            out.append(buf.getvalue())
        except Exception as exc:
            logger.warning("pdfium render page %d failed: %s", index, exc)
            continue
    return out
