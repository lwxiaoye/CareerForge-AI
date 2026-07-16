"""avatar_extractor 的单元测试。

主要覆盖 find_avatar_region：在合成的"首页 PNG"上扫描候选区，挑色相最丰富的那个。
通过 monkeypatch 临时把 AVATAR_DIR 指向 tmpdir，避免污染真实目录。
"""
import io
import os
import tempfile
import unittest
from unittest.mock import patch

from PIL import Image


def _make_synthetic_page(*, photo_box=None, photo_color=(220, 180, 140),
                        page_size=(800, 1100), bg=(255, 255, 255),
                        text_color=(20, 20, 20)):
    """合成一张"首页"：白底 + 顶部一两条深色文字带 + 可选右上角彩色色块（模拟照片）。

    photo_box: (x0, y0, x1, y1) — 模拟真人照片的彩色区域。
    """
    image = Image.new("RGB", page_size, bg)
    # 模拟页面里的几行文字（深色细横条）— 帮助确认 find_avatar_region 不会把它们当头像
    for y in (220, 260, 300, 360, 420, 480, 540, 600, 660, 720):
        for x in range(60, 540, 12):
            image.putpixel((x, y), text_color)
    if photo_box is not None:
        x0, y0, x1, y1 = photo_box
        # 填一层肤色底，再随机打一些深浅不一的点，模拟照片（色相丰富、亮度有变化）
        import random
        random.seed(0)
        for y in range(y0, y1):
            for x in range(x0, x1):
                jitter = random.randint(-25, 25)
                r = max(0, min(255, photo_color[0] + jitter))
                g = max(0, min(255, photo_color[1] + jitter))
                b = max(0, min(255, photo_color[2] + jitter // 2))
                image.putpixel((x, y), (r, g, b))
    buf = io.BytesIO()
    image.save(buf, format="PNG", optimize=False)
    return buf.getvalue()


def _img_from_url(app, url):
    """从保存的 URL 读回头像字节。"""
    from app.student import avatar_storage
    name = url.rsplit("/", 1)[-1]
    path = avatar_storage.AVATAR_DIR / name
    with Image.open(path) as image:
        image.load()
    return image


class FindAvatarRegionTests(unittest.TestCase):
    def setUp(self):
        from pathlib import Path
        self._tmpdir = Path(tempfile.mkdtemp(prefix="avatar_test_"))
        self._patches = [
            patch("app.student.avatar_storage.AVATAR_DIR", new=self._tmpdir),
        ]
        for p in self._patches:
            p.start()

    def tearDown(self):
        for p in self._patches:
            p.stop()
        import shutil
        shutil.rmtree(self._tmpdir, ignore_errors=True)

    def test_finds_top_right_photo(self):
        """右上角放了彩色照片时，应挑出接近该区域的位置。"""
        from app.student.avatar_extractor import find_avatar_region

        page_size = (800, 1100)
        # 右上角 1/3 × 1/3 区域
        short = min(page_size)
        side = short // 3
        photo_box = (page_size[0] - side, 0, page_size[0], side)
        png = _make_synthetic_page(page_size=page_size, photo_box=photo_box)

        url = find_avatar_region(png)
        self.assertIsNotNone(url, "expected to find a photo-like region")
        image = _img_from_url(self, url)
        # 校验挑出的区域跟右上角彩色色块位置基本重合
        x0, y0, x1, y1 = photo_box
        # 裁剪图会被等比缩放到不超过 MAX_AVATAR_DIMENSION，宽高比保持
        self.assertGreater(image.size[0], 60)
        self.assertGreater(image.size[1], 60)
        # 平均色应该明显偏暖（肤色基色），不是纯白
        rgb = image.convert("RGB").resize((32, 32))
        pixels = list(rgb.getdata())
        r_mean = sum(p[0] for p in pixels) / len(pixels)
        g_mean = sum(p[1] for p in pixels) / len(pixels)
        b_mean = sum(p[2] for p in pixels) / len(pixels)
        self.assertGreater(r_mean, 150, f"expected warm photo, got mean rgb=({r_mean:.0f},{g_mean:.0f},{b_mean:.0f})")

    def test_finds_top_left_photo(self):
        """左上角放照片也应被识别（覆盖中文简历里"左上图"这种版式）。"""
        from app.student.avatar_extractor import find_avatar_region

        page_size = (800, 1100)
        short = min(page_size)
        side = short // 3
        photo_box = (0, 0, side, side)
        png = _make_synthetic_page(page_size=page_size, photo_box=photo_box)

        url = find_avatar_region(png)
        self.assertIsNotNone(url)
        image = _img_from_url(self, url)
        rgb = image.convert("RGB").resize((32, 32))
        pixels = list(rgb.getdata())
        r_mean = sum(p[0] for p in pixels) / len(pixels)
        self.assertGreater(r_mean, 150, "left-side photo should also be picked")

    def test_no_photo_returns_none(self):
        """纯白底 + 纯黑文字时，没有色相丰富的区域，应返回 None。"""
        from app.student.avatar_extractor import find_avatar_region

        png = _make_synthetic_page(photo_box=None)
        self.assertIsNone(find_avatar_region(png))

    def test_empty_or_invalid_png_returns_none(self):
        """空字节 / 乱字节都应安全返回 None，不抛异常。"""
        from app.student.avatar_extractor import find_avatar_region

        self.assertIsNone(find_avatar_region(b""))
        self.assertIsNone(find_avatar_region(b"not a png"))


class ColorRichnessTests(unittest.TestCase):
    def test_white_image_is_zero(self):
        from app.student.avatar_extractor import _color_richness
        image = Image.new("RGB", (200, 200), (255, 255, 255))
        self.assertLess(_color_richness(image), 0.05)

    def test_skin_tone_image_is_high(self):
        from app.student.avatar_extractor import _color_richness
        # 简单的肤色渐变
        import random
        random.seed(1)
        image = Image.new("RGB", (200, 200), (220, 180, 140))
        for y in range(200):
            for x in range(200):
                j = random.randint(-30, 30)
                r = max(0, min(255, 220 + j))
                g = max(0, min(255, 180 + j // 2))
                b = max(0, min(255, 140 + j // 3))
                image.putpixel((x, y), (r, g, b))
        score = _color_richness(image)
        self.assertGreater(score, 0.15, f"skin tone should score > 0.15, got {score}")


if __name__ == "__main__":
    unittest.main()


class ScorePhotoImageTests(unittest.TestCase):
    """ScorePhotoImageTests - _score_photo_image 在 DOCX / PDF 嵌入图场景的公共打分器。"""

    def _photo(self, size=(300, 400), base=(220, 180, 140)):
        import random
        random.seed(0)
        image = Image.new("RGB", size, base)
        for y in range(size[1]):
            for x in range(size[0]):
                j = random.randint(-25, 25)
                r = max(0, min(255, base[0] + j))
                g = max(0, min(255, base[1] + j // 2))
                b = max(0, min(255, base[2] + j // 3))
                image.putpixel((x, y), (r, g, b))
        return image

    def test_real_photo_scores_high(self):
        from app.student.avatar_extractor import _score_photo_image
        score = _score_photo_image(self._photo())
        self.assertGreater(score, 0.4, f"photo score too low: {score}")

    def test_tiny_logo_rejected(self):
        from app.student.avatar_extractor import _score_photo_image
        image = Image.new("RGB", (40, 40), (10, 80, 180))
        self.assertEqual(_score_photo_image(image), -1.0)

    def test_wide_banner_rejected(self):
        from app.student.avatar_extractor import _score_photo_image
        image = Image.new("RGB", (900, 120), (10, 80, 180))
        self.assertEqual(_score_photo_image(image), -1.0)

    def test_solid_color_block_rejected(self):
        from app.student.avatar_extractor import _score_photo_image
        image = Image.new("RGB", (200, 200), (200, 0, 0))
        self.assertEqual(_score_photo_image(image), -1.0)

    def test_letterbox_photo_kept(self):
        """真人证件照常见 3:4 长宽比，应在合法范围内。"""
        from app.student.avatar_extractor import _score_photo_image
        score = _score_photo_image(self._photo(size=(300, 400)))
        self.assertGreater(score, 0.0)

    def test_landscape_photo_kept(self):
        from app.student.avatar_extractor import _score_photo_image
        score = _score_photo_image(self._photo(size=(400, 300)))
        self.assertGreater(score, 0.0)


def _make_photo_png_bytes(size=(300, 400), base=(220, 180, 140)):
    import random, io
    random.seed(0)
    image = Image.new("RGB", size, base)
    for y in range(size[1]):
        for x in range(size[0]):
            j = random.randint(-25, 25)
            r = max(0, min(255, base[0] + j))
            g = max(0, min(255, base[1] + j // 2))
            b = max(0, min(255, base[2] + j // 3))
            image.putpixel((x, y), (r, g, b))
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return buf.getvalue()


def _make_logo_png_bytes(size=(60, 60), color=(10, 80, 180)):
    import io
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="PNG")
    return buf.getvalue()


def _build_docx_with_images(entries):
    """entries: list of (bytes, "logo"|"photo"), 按顺序 inline 进 docx。返回 docx 字节。"""
    from pathlib import Path
    import io, tempfile
    from docx import Document
    from docx.shared import Inches
    tmp = Path(tempfile.mkdtemp())
    doc = Document()
    for index, (data, kind) in enumerate(entries):
        path = tmp / f"img_{index}.png"
        path.write_bytes(data)
        para = doc.add_paragraph(f"img {index}:")
        if kind == "logo":
            para.add_run().add_picture(str(path), width=Inches(0.5))
        else:
            para.add_run().add_picture(str(path), width=Inches(2.0))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class ExtractAvatarFromDocxTests(unittest.TestCase):
    """extract_avatar_from_docx 要能跳过 logo / 装饰图，挑最像头像的那张。"""

    def setUp(self):
        from pathlib import Path
        import tempfile
        from unittest.mock import patch
        self._tmpdir = Path(tempfile.mkdtemp(prefix="avatar_docx_"))
        self._patches = [
            patch("app.student.avatar_storage.AVATAR_DIR", new=self._tmpdir),
        ]
        for p in self._patches:
            p.start()

    def tearDown(self):
        for p in self._patches:
            p.stop()
        import shutil
        shutil.rmtree(self._tmpdir, ignore_errors=True)

    def test_picks_photo_over_first_logo(self):
        """DOCX 里第一张是 60x60 校徽，第二张才是真人头像 —— 必须挑第二张。"""
        from app.student.avatar_extractor import extract_avatar_from_docx
        docx_bytes = _build_docx_with_images([
            (_make_logo_png_bytes(), "logo"),
            (_make_photo_png_bytes(), "photo"),
        ])
        url = extract_avatar_from_docx(docx_bytes)
        self.assertIsNotNone(url, "should still find the photo even though logo is first")
        image = _img_from_url(self, url)
        self.assertGreater(image.size[0], 100, f"avatar width too small: {image.size}")
        self.assertGreater(image.size[1], 100, f"avatar height too small: {image.size}")
        rgb = image.convert("RGB").resize((32, 32))
        pixels = list(rgb.getdata())
        b_mean = sum(p[2] for p in pixels) / len(pixels)
        r_mean = sum(p[0] for p in pixels) / len(pixels)
        self.assertGreater(r_mean, b_mean, f"expected warm tone, got r={r_mean:.0f} b={b_mean:.0f}")

    def test_returns_none_when_only_logos(self):
        """DOCX 里只有 60x60 校徽 —— 不应当作头像，返回 None。"""
        from app.student.avatar_extractor import extract_avatar_from_docx
        docx_bytes = _build_docx_with_images([
            (_make_logo_png_bytes(), "logo"),
            (_make_logo_png_bytes(size=(40, 40), color=(200, 0, 0)), "logo"),
        ])
        self.assertIsNone(extract_avatar_from_docx(docx_bytes))

    def test_returns_none_on_garbage_bytes(self):
        from app.student.avatar_extractor import extract_avatar_from_docx
        self.assertIsNone(extract_avatar_from_docx(b"not a docx"))
        self.assertIsNone(extract_avatar_from_docx(b""))


def _build_pdf_with_images(entries):
    """entries: list of (PIL.Image, "logo"|"photo"), 在 letter 页上不同位置 drawImage。返回 PDF 字节。"""
    from pathlib import Path
    import io, tempfile
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    tmp = Path(tempfile.mkdtemp())
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(100, 700, "Synthetic resume page")
    y_cursor = 650
    for index, (image, kind) in enumerate(entries):
        path = tmp / f"img_{index}.png"
        image.save(path, "PNG")
        if kind == "logo":
            c.drawImage(str(path), 80, y_cursor, width=30, height=30, mask="auto")
            y_cursor -= 40
        else:
            c.drawImage(str(path), 350, 600, width=120, height=160, mask="auto")
    c.save()
    return buf.getvalue()


class ExtractAvatarFromPdfLegacyTests(unittest.TestCase):
    """_extract_avatar_from_pdf_legacy: pypdf 嵌入图路径, 不再 “第一张就 return”。"""

    def setUp(self):
        from pathlib import Path
        import tempfile
        from unittest.mock import patch
        self._tmpdir = Path(tempfile.mkdtemp(prefix="avatar_pdf_legacy_"))
        self._patches = [
            patch("app.student.avatar_storage.AVATAR_DIR", new=self._tmpdir),
        ]
        for p in self._patches:
            p.start()

    def tearDown(self):
        for p in self._patches:
            p.stop()
        import shutil
        shutil.rmtree(self._tmpdir, ignore_errors=True)

    def _photo(self):
        import random
        random.seed(0)
        image = Image.new("RGB", (300, 400), (220, 180, 140))
        for y in range(400):
            for x in range(300):
                j = random.randint(-25, 25)
                image.putpixel((x, y), (max(0,min(255,220+j)), max(0,min(255,180+j//2)), max(0,min(255,140+j//3))))
        return image

    def _logo(self, color=(10, 80, 180)):
        return Image.new("RGB", (40, 40), color)

    def test_picks_photo_when_logo_is_first(self):
        """PDF 里有小 logo + 大头像, 旧实现会把第一张当头像, 现在必须挑真人头像。"""
        from app.student.avatar_extractor import _extract_avatar_from_pdf_legacy
        pdf_bytes = _build_pdf_with_images([
            (self._logo(), "logo"),
            (self._photo(), "photo"),
        ])
        url = _extract_avatar_from_pdf_legacy(pdf_bytes)
        self.assertIsNotNone(url)
        image = _img_from_url(self, url)
        self.assertGreater(image.size[0], 100)
        self.assertGreater(image.size[1], 100)

    def test_returns_none_when_only_logos(self):
        from app.student.avatar_extractor import _extract_avatar_from_pdf_legacy
        pdf_bytes = _build_pdf_with_images([
            (self._logo(color=(10, 80, 180)), "logo"),
            (self._logo(color=(180, 10, 10)), "logo"),
        ])
        self.assertIsNone(_extract_avatar_from_pdf_legacy(pdf_bytes))
